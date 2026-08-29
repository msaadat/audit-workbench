"""Build the Meridian Bank treasury pack's documents: three criteria documents
as .docx, and one PDF deal pack per sampled deal.

Every figure printed here is read from ground_truth.json, so the paper and the
populations agree except where an exception deliberately makes them disagree.
"""

from __future__ import annotations

import json
from datetime import date, datetime, timedelta
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (PageBreak, Paragraph, SimpleDocTemplate, Spacer,
                                Table, TableStyle)

ROOT = Path(__file__).resolve().parent.parent / "documents"
PACKS_DIR = ROOT / "deal-packs"
GT = json.loads((Path(__file__).resolve().parent / "ground_truth.json").read_text())

DEALS = GT["deals"]
CONFS = GT["confirmations"]
SETTS = GT["settlements"]
STAFF = GT["staff"]
CPS = GT["counterparties"]
BROKERS = GT["brokers"]
SSIS = {s["SSI_ID"]: s for s in GT["ssis"]}
PACKS = GT["packs"]
EXTRA = GT["pack_extras"]

ENTITY = "Meridian Bank Limited"
ENTITY_SWIFT = "MRDBPKKA"

CP_SWIFT = {
    "CP-001": "STUBPKKA", "CP-002": "HRBNPKKA", "CP-003": "PNCLPKLA",
    "CP-004": "HRZNPKKA", "CP-005": "CRINPKKA", "CP-006": "APXIPKKA",
    "CP-007": "FRDBPKIS", "CP-008": "NRGTPKLA", "CP-009": "CEDCPKKA",
    "CP-010": "SBPPPKKA",
}
CP_CITY = {
    "CP-001": "Karachi", "CP-002": "Karachi", "CP-003": "Lahore",
    "CP-004": "Karachi", "CP-005": "Karachi", "CP-006": "Karachi",
    "CP-007": "Islamabad", "CP-008": "Lahore", "CP-009": "Karachi",
    "CP-010": "Karachi",
}
PRODUCT = {
    "FX_SPOT": "Foreign exchange - spot",
    "FX_FORWARD": "Foreign exchange - forward",
    "MM_PLACEMENT": "Money market - placement",
    "MM_BORROWING": "Money market - borrowing",
    "TBILL_PURCHASE": "Government securities - treasury bill purchase",
}
CONF_TITLE = {
    "SWIFT MT300": "CONFIRMATION OF FOREIGN EXCHANGE TRANSACTION (MT300)",
    "SWIFT MT320": "CONFIRMATION OF FIXED TERM DEPOSIT / LOAN (MT320)",
    "SWIFT MT518": "CONFIRMATION OF SECURITIES TRANSACTION (MT518)",
}

HOLIDAYS = {date(2025, 1, 1), date(2025, 2, 5), date(2025, 3, 31),
            date(2025, 4, 1), date(2025, 4, 2), date(2025, 5, 1),
            date(2025, 6, 6), date(2025, 6, 9), date(2025, 6, 10)}


def add_bd(d: date, n: int) -> date:
    cur = d
    while n > 0:
        cur += timedelta(days=1)
        if cur.weekday() < 5 and cur not in HOLIDAYS:
            n -= 1
    return cur


def long_date(value: str) -> str:
    d = date.fromisoformat(value)
    return f"{d.day} {d.strftime('%B %Y')}"


def money(value, currency: str = "PKR") -> str:
    return f"{currency} {float(value):,.2f}"


# --------------------------------------------------------------------------
# criteria documents
# --------------------------------------------------------------------------

def docx_base() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    return doc


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def build_policy() -> None:
    doc = docx_base()
    title = doc.add_paragraph()
    run = title.add_run(f"{ENTITY.upper()}\nTREASURY AND INVESTMENT POLICY")
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.add_run("Extract for internal audit - sections 4 to 8\n").italic = True
    sub.add_run("Document reference TP/2023/04  |  Version 4.1  |  Approved by "
                "the Board Risk Committee on 12 October 2023  |  Effective 1 "
                "November 2023  |  Next review due 1 November 2024").italic = True
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    sections = [
        ("4. Organisation and segregation of duties", [
            ("4.1", "Dealing, confirmation and settlement are performed by "
             "separate functions. A member of the front office shall not "
             "confirm, verify or release settlement of any deal, whether or "
             "not that deal is one they executed."),
            ("4.2", "The middle office monitors counterparty and dealer limits "
             "independently of the dealing room and reports utilisation daily "
             "to the Head of Treasury and the Head of Risk."),
            ("4.3", "Every settlement instruction is released by one officer "
             "and approved by a second. Instructions for an amount at or above "
             "the threshold in clause 8.5 require a further approver."),
        ]),
        ("5. Dealing", [
            ("5.1", "Deals are executed only with counterparties carrying a "
             "current approved limit and a status of Active on the "
             "counterparty master. A counterparty whose limit has expired, or "
             "whose status is Suspended or Under Review, is closed to new "
             "business until the position is restored by the Board Risk "
             "Committee."),
            ("5.2", "Dealing is permitted between 09:00 and 17:00 on business "
             "days only. Any deal struck outside those hours, or on a "
             "non-business day, requires the prior written approval of the "
             "Head of Treasury and must be reported to ALCO at its next "
             "meeting."),
            ("5.3", "A dealer may transact only the products named in their "
             "dealing authorisation, and only while that authorisation is "
             "current. No deal may exceed the dealer's per-deal limit, and the "
             "aggregate value a dealer transacts on one day may not exceed "
             "their daily limit."),
            ("5.4", "Every deal is captured in the treasury system within 15 "
             "minutes of execution. Capture records the dealer's identity and "
             "the time of execution as distinct from the time of capture."),
            ("5.5", "The dealt rate is evidenced against the market rate "
             "prevailing at the time of execution. Deals struck outside the "
             "tolerance in clause 6.3 are escalated to the Head of Treasury "
             "before confirmation."),
            ("5.6", "A deal may not be amended or cancelled after capture "
             "without the approval of the Head of Treasury Operations. The "
             "approver's identity and the date of approval are recorded "
             "against the deal."),
        ]),
        ("6. Rates and valuation", [
            ("6.1", "The middle office publishes an independent market rate "
             "file each business day covering every instrument and tenor the "
             "bank deals in, together with the day's traded high and low."),
            ("6.2", "The rate file is the reference for rate reasonableness "
             "testing. It is not sourced from the dealing room."),
            ("6.3", "A foreign exchange deal shall be struck within 25 basis "
             "points of the published mid rate for its instrument and tenor. A "
             "money market or treasury bill deal shall be struck within 15 "
             "basis points. Deviations beyond these tolerances require prior "
             "approval and are reported to ALCO."),
        ]),
        ("7. Confirmation", [
            ("7.1", "Every deal is confirmed with the counterparty. The "
             "confirmation is issued by the back office from the captured "
             "deal record, not from the dealer's ticket."),
            ("7.2", "Confirmations are despatched no later than one business "
             "day after the deal date."),
            ("7.3", "The counterparty's confirmation is matched to the deal "
             "record on counterparty, principal, rate, value date and maturity "
             "date. A discrepancy is referred to the Head of Treasury "
             "Operations and is resolved before settlement."),
            ("7.4", "No deal is settled before the counterparty's confirmation "
             "has been received and matched. A confirmation produced "
             "internally is not evidence of the counterparty's agreement."),
        ]),
        ("8. Settlement", [
            ("8.1", "Settlement is effected on the contracted value date. A "
             "settlement made early or late is an exception and is reported."),
            ("8.2", "The settlement amount is computed from the confirmed "
             "principal and the confirmed rate and is agreed to the deal "
             "record before release."),
            ("8.3", "Funds are remitted only to an account held on the "
             "approved standing settlement instruction file for that "
             "counterparty and currency. Ad hoc payment instructions are not "
             "accepted."),
            ("8.4", "An amendment to a standing settlement instruction takes "
             "effect two business days after approval. No settlement may be "
             "made to an amended instruction inside that period."),
            ("8.5", "A settlement instruction for PKR 250,000,000 or more "
             "requires two authorised signatures on the instruction itself in "
             "addition to system release and approval."),
            ("8.6", "The nostro reconciliation officer reconciles every "
             "settlement to the account statement within two business days."),
        ]),
    ]
    for name, clauses in sections:
        heading(doc, name, 1)
        for number, text in clauses:
            para = doc.add_paragraph()
            para.add_run(f"{number}  ").bold = True
            para.add_run(text)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("This extract is issued for the purposes of the internal audit "
                 "of treasury operations for the half year ended 30 June 2025. "
                 "It reproduces sections 4 to 8 of the policy in full. Sections "
                 "1 to 3 (scope, definitions and governance) and 9 to 11 "
                 "(reporting, breaches and review) are not reproduced."
                 ).italic = True
    doc.save(ROOT / "01_Treasury_and_Investment_Policy.docx")


def build_limits() -> None:
    doc = docx_base()
    title = doc.add_paragraph()
    run = title.add_run(f"{ENTITY.upper()}\nCOUNTERPARTY AND DEALER LIMIT MATRIX")
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph(
        "The limits below govern treasury dealing. Counterparty limits are "
        "expressed as the maximum aggregate principal that may be outstanding "
        "with a counterparty at any time. Dealer limits are expressed per deal "
        "and per dealing day.")

    heading(doc, "Counterparty limits", 1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells,
                          ("Counterparty", "Type", "Rating", "Limit (PKR)",
                           "Limit expiry", "Status")):
        cell.paragraphs[0].add_run(text).bold = True
    for cp_id, cp in CPS.items():
        row = table.add_row().cells
        row[0].text = f"{cp['name']} ({cp_id})"
        row[1].text = cp["type"]
        row[2].text = cp["rating"]
        row[3].text = f"{cp['limit']:,}"
        row[4].text = ""
        row[5].text = cp["status"]
    doc.add_paragraph()
    doc.add_paragraph(
        "Limit expiry dates and status effective dates are maintained on the "
        "counterparty master in the treasury system and are not reproduced "
        "here.").italic = True

    heading(doc, "Dealer limits", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells,
                          ("Dealer", "Role", "Per deal (PKR)", "Per day (PKR)",
                           "Authorised products")):
        cell.paragraphs[0].add_run(text).bold = True
    label = {"FX_SPOT": "FX spot", "FX_FORWARD": "FX forward",
             "MM_PLACEMENT": "money market placement",
             "MM_BORROWING": "money market borrowing",
             "TBILL_PURCHASE": "treasury bills"}
    with (ROOT.parent / "data" / "02_dealer_limits.csv").open() as fh:
        import csv as _csv
        for r in _csv.DictReader(fh):
            dealer = r["DEALER_ID"]
            row = table.add_row().cells
            row[0].text = f"{STAFF[dealer]['name']} ({dealer})"
            row[1].text = STAFF[dealer]["title"]
            row[2].text = f'{int(r["PER_DEAL_LIMIT_PKR"]):,}'
            row[3].text = f'{int(r["DAILY_LIMIT_PKR"]):,}'
            products = ", ".join(
                label[p] for p in r["AUTHORISED_PRODUCTS"].split(";"))
            row[4].text = products[0].upper() + products[1:]

    doc.add_paragraph()
    doc.add_paragraph(
        "Dealing authorisations are granted individually and lapse on the date "
        "recorded against each dealer in the treasury system.")
    doc.save(ROOT / "02_Counterparty_and_Dealer_Limit_Matrix.docx")


def build_minutes() -> None:
    doc = docx_base()
    title = doc.add_paragraph()
    run = title.add_run(f"{ENTITY.upper()}\nMINUTES OF MEETING")
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.add_run("Internal audit planning - treasury operations").italic = True
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for label, value in (
            ("Date", "23 July 2025"),
            ("Venue", "Board room, head office, Karachi"),
            ("Present", "Nadia Rahman, Chief Financial Officer (chair); "
                        "Imran Qadir, Head of Treasury; Nida Hasan, Head of "
                        "Risk; Rehan Baig, Head of Treasury Operations; "
                        "Lubna Farooq, Compliance Officer; Sarah Idrees, Head "
                        "of Internal Audit"),
            ("Subject", "Scope of the internal audit of treasury operations "
                        "for the half year ended 30 June 2025")):
        para = doc.add_paragraph()
        para.add_run(f"{label}: ").bold = True
        para.add_run(value)

    items = [
        ("1. Background", [
            "The Head of Treasury reported that dealing volumes had grown "
            "materially over the half year, driven by the bank's liquidity "
            "position and by wider spreads in the interbank market. One "
            "thousand deals were struck in the six months to 30 June 2025 "
            "across foreign exchange, money market and treasury bills.",
            "The Chief Financial Officer noted that the dealing room had "
            "operated through the period with one dealer fewer than "
            "establishment, and that a junior dealer had been on secondment "
            "from June.",
        ]),
        ("2. Control environment", [
            "The Head of Treasury Operations reported that confirmations had "
            "on occasion been despatched late where the counterparty's own "
            "back office was slow to respond, and that a small number of "
            "deals had been amended after capture.",
            "The Head of Risk observed that dealer authorisations had not "
            "been reviewed since the start of the financial year, and that at "
            "least one dealer was operating without a current limit record. "
            "She undertook to have the limit file reconciled to the "
            "establishment before the audit fieldwork began. As at the date of "
            "these minutes that reconciliation had not been completed.",
            "The Compliance Officer reported that the counterparty master "
            "carried two counterparties whose position had changed during the "
            "period, one suspended and one whose limit had lapsed, and that "
            "she had not received confirmation that dealing with either had "
            "stopped.",
        ]),
        ("3. Matters raised by management", [
            "The Head of Treasury stated that the dealing room regards the "
            "rate tolerance in clause 6.3 of the policy as guidance rather "
            "than a hard limit in fast markets, and that escalation of "
            "out-of-tolerance deals had not in practice occurred.",
            "The Head of Treasury Operations stated that settlement had on "
            "occasion been released against a confirmation that had not yet "
            "been matched, where the value date fell on the day of dealing.",
            "No exception report covering deals, confirmations or settlements "
            "has been produced for the period.",
        ]),
        ("4. Audit scope agreed", [
            "Internal Audit will perform a risk-based review of treasury "
            "dealing, confirmation and settlement for the half year ended 30 "
            "June 2025, covering dealer and counterparty limit compliance, "
            "rate reasonableness, segregation of duties, confirmation and "
            "settlement timeliness and completeness, and settlement routing.",
            "Treasury will provide the deal, confirmation and settlement "
            "populations, the counterparty and dealer limit files, the "
            "independent market rate file, the standing settlement instruction "
            "file, and supporting documentation for a sample of deals.",
            "The Head of Internal Audit noted that the Treasury and Investment "
            "Policy carries a review date of 1 November 2024 which has passed, "
            "and asked that the current approved version be confirmed. The "
            "Chief Financial Officer undertook to revert.",
        ]),
    ]
    for name, paras in items:
        heading(doc, name, 1)
        for text in paras:
            doc.add_paragraph(text)

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.add_run("There being no other business the meeting closed at 11:40. "
                    "Minutes taken by the Compliance Officer and circulated for "
                    "comment; not yet signed as approved.").italic = True
    doc.save(ROOT / "03_Minutes_Internal_Audit_Planning.docx")


# --------------------------------------------------------------------------
# deal packs
# --------------------------------------------------------------------------

styles = getSampleStyleSheet()
S_ORG = ParagraphStyle("org", parent=styles["Normal"], fontName="Helvetica-Bold",
                       fontSize=13, leading=16, alignment=TA_CENTER)
S_ORGSUB = ParagraphStyle("orgsub", parent=styles["Normal"], fontSize=8.5,
                          leading=11, alignment=TA_CENTER,
                          textColor=colors.HexColor("#444444"))
S_TITLE = ParagraphStyle("title", parent=styles["Normal"],
                         fontName="Helvetica-Bold", fontSize=11, leading=14,
                         alignment=TA_CENTER, spaceBefore=10, spaceAfter=8)
S_BODY = ParagraphStyle("body", parent=styles["Normal"], fontSize=9, leading=12)
S_SMALL = ParagraphStyle("small", parent=styles["Normal"], fontSize=7.5,
                         leading=10, textColor=colors.HexColor("#555555"))
S_MONO = ParagraphStyle("mono", parent=styles["Normal"], fontName="Courier",
                        fontSize=8.5, leading=11)
S_HAND = ParagraphStyle("hand", parent=styles["Normal"],
                        fontName="Courier-Oblique", fontSize=9, leading=12)

GRID = TableStyle([
    ("FONT", (0, 0), (-1, -1), "Helvetica", 9),
    ("FONT", (0, 0), (0, -1), "Helvetica-Bold", 9),
    ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#333333")),
    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ("TOPPADDING", (0, 0), (-1, -1), 3),
    ("LINEBELOW", (0, 0), (-1, -2), 0.25, colors.HexColor("#DDDDDD")),
])


def kv(rows, widths=(52 * mm, 116 * mm)) -> Table:
    body = [[k, v if isinstance(v, Paragraph) else Paragraph(str(v), S_BODY)]
            for k, v in rows]
    table = Table(body, colWidths=widths)
    table.setStyle(GRID)
    return table


def rule(width=168 * mm):
    table = Table([[""]], colWidths=(width,), rowHeights=(1,))
    table.setStyle(TableStyle([("LINEABOVE", (0, 0), (-1, 0), 0.8,
                               colors.HexColor("#222222"))]))
    return table


def letterhead(name: str, lines: list[str]) -> list:
    out = [Paragraph(name.upper(), S_ORG)]
    for line in lines:
        out.append(Paragraph(line, S_ORGSUB))
    out += [Spacer(1, 4), rule(), Spacer(1, 2)]
    return out


def signatures(entries: list[tuple[str, str, str]]) -> Table:
    """entries: (role, name, dated) - an empty name prints an unsigned block."""
    cells, widths = [], []
    for role, name, dated in entries:
        block = [Paragraph(f"<b>{role}</b>", S_SMALL), Spacer(1, 16)]
        block.append(Paragraph("_" * 26, S_BODY))
        block.append(Paragraph(name or "&nbsp;", S_SMALL))
        block.append(Paragraph(dated or "&nbsp;", S_SMALL))
        cells.append(block)
        widths.append(168 * mm / len(entries))
    table = Table([cells], colWidths=widths)
    table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"),
                               ("LEFTPADDING", (0, 0), (-1, -1), 0)]))
    return table


def person(staff_id: str) -> str:
    s = STAFF[staff_id]
    return f"{s['name']} ({staff_id}), {s['title']}"


# --------------------------------------------------------------------------
# the sampled deal that the system never recorded
# --------------------------------------------------------------------------

SYNTHETIC = {
    "DEAL_ID": "TD-2025-0447",
    "DEAL_DATE": "2025-03-19",
    "DEAL_TIME": "14:52:31",
    "DEAL_TYPE": "MM_PLACEMENT",
    "INSTRUMENT": "PKR_MM",
    "TENOR": "1M",
    "COUNTERPARTY_ID": "CP-003",
    "DEALER_ID": "TS-007",
    "BROKER_ID": "BR-001",
    "CURRENCY": "PKR",
    "NOTIONAL_AMOUNT": "175000000.00",
    "DEAL_RATE": "11.4820",
    "NOTIONAL_PKR": "175000000.00",
    "VALUE_DATE": "2025-03-20",
    "MATURITY_DATE": "2025-04-21",
    "DEAL_STATUS": "Settled",
    "CAPTURED_BY_ID": "TS-007",
    "CAPTURE_TIMESTAMP": "2025-03-19 14:58:06",
    "AMENDED_FLAG": "N",
    "AMENDMENT_DATE": "",
    "AMENDMENT_APPROVED_BY_ID": "",
}
SYNTHETIC_CONF = {
    "CONFIRMATION_ID": "CNF-2025-M041",
    "DEAL_ID": "TD-2025-0447",
    "CONFIRMATION_MODE": "SWIFT MT320",
    "SENT_DATE": "2025-03-19",
    "COUNTERPARTY_RECEIVED_DATE": "2025-03-20",
    "CONFIRMED_BY_ID": "TS-015",
    "MATCH_STATUS": "Matched",
    "DISCREPANCY_NOTE": "",
}
SYNTHETIC_STL = {
    "SETTLEMENT_ID": "STL-2025-M014",
    "DEAL_ID": "TD-2025-0447",
    "SETTLEMENT_DATE": "2025-03-20",
    "SETTLEMENT_CURRENCY": "PKR",
    "SETTLEMENT_AMOUNT": "175000000.00",
    "SETTLEMENT_AMOUNT_PKR": "175000000.00",
    "SSI_ID": next((s["SSI_ID"] for s in GT["ssis"]
                    if s["COUNTERPARTY_ID"] == "CP-003" and s["CURRENCY"] == "PKR"), ""),
    "BENEFICIARY_BANK": "Pinnacle Commercial Bank, Lahore",
    "BENEFICIARY_ACCOUNT": next((s["BENEFICIARY_ACCOUNT"] for s in GT["ssis"]
                                 if s["COUNTERPARTY_ID"] == "CP-003"
                                 and s["CURRENCY"] == "PKR"), ""),
    "NOSTRO_ACCOUNT": "SBP-CA-0041-MBL",
    "RELEASED_BY_ID": "TS-017",
    "APPROVED_BY_ID": "TS-014",
    "SECOND_APPROVER_ID": "",
    "PAYMENT_REFERENCE": "PMT-2025-M0114",
}


# --------------------------------------------------------------------------
# pages
# --------------------------------------------------------------------------

def page_deal_ticket(deal: dict, flags: set[str]) -> list:
    cp = CPS[deal["COUNTERPARTY_ID"]]
    fx = deal["DEAL_TYPE"] in ("FX_SPOT", "FX_FORWARD")
    story = letterhead(ENTITY, ["Treasury - dealing room, head office, Karachi",
                                f"SWIFT {ENTITY_SWIFT}"])
    story.append(Paragraph("DEALING TICKET", S_TITLE))

    if "altered_rate" in flags:
        original = float(deal["DEAL_RATE"]) * 0.9938
        rate_cell = Paragraph(
            f'<strike>{original:.4f}</strike> &nbsp;&nbsp;'
            f'<font face="Courier-Oblique">{float(deal["DEAL_RATE"]):.4f}</font>'
            f' &nbsp;<font size="7" color="#666666">(amended on the ticket in '
            f'ink; not initialled)</font>', S_BODY)
    else:
        rate_cell = Paragraph(
            f'{float(deal["DEAL_RATE"]):.4f}'
            f'{" per cent per annum" if not fx else ""}', S_BODY)

    rows = [
        ("Deal reference", deal["DEAL_ID"]),
        ("Deal date", f'{long_date(deal["DEAL_DATE"])} at {deal["DEAL_TIME"]}'),
        ("Product", PRODUCT[deal["DEAL_TYPE"]]),
        ("Instrument / tenor", f'{deal["INSTRUMENT"]} / {deal["TENOR"]}'),
        ("Counterparty", f'{cp["name"]} ({deal["COUNTERPARTY_ID"]}), '
                         f'{CP_CITY[deal["COUNTERPARTY_ID"]]}'),
        ("Broker", BROKERS[deal["BROKER_ID"]] if deal["BROKER_ID"] else "Direct"),
        ("Principal", money(deal["NOTIONAL_AMOUNT"], deal["CURRENCY"])),
    ]
    if fx:
        rows.append(("PKR equivalent", money(deal["NOTIONAL_PKR"])))
    rows += [
        ("Rate", rate_cell),
        ("Value date", long_date(deal["VALUE_DATE"])),
        ("Maturity date", long_date(deal["MATURITY_DATE"])),
        ("Captured in treasury system",
         f'{deal["CAPTURE_TIMESTAMP"]} by {deal["CAPTURED_BY_ID"]}'),
    ]
    story.append(kv(rows))
    story.append(Spacer(1, 14))

    dealer = person(deal["DEALER_ID"])
    if "no_authorisation" in flags:
        auth = ("Authorised by (Head of Treasury or Chief Dealer)", "", "")
    elif "late_authorisation" in flags:
        auth_date = add_bd(date.fromisoformat(deal["DEAL_DATE"]), 2)
        auth = ("Authorised by (Head of Treasury or Chief Dealer)",
                person("TS-002"), f"Dated {long_date(auth_date.isoformat())}")
    else:
        supervisor = "TS-002" if deal["DEALER_ID"] != "TS-002" else "TS-001"
        auth = ("Authorised by (Head of Treasury or Chief Dealer)",
                person(supervisor), f'Dated {long_date(deal["DEAL_DATE"])}')
    story.append(signatures([
        ("Dealt by", dealer, f'Dated {long_date(deal["DEAL_DATE"])}'), auth]))
    story.append(Spacer(1, 10))
    story.append(Paragraph(
        "Form TR-01. Retain with the counterparty confirmation and the "
        "settlement instruction. Fictional document produced for audit "
        "training.", S_SMALL))
    return story


def broker_note_ref(deal: dict) -> str:
    return (f'{deal["BROKER_ID"]}-{deal["DEAL_DATE"][:4]}-'
            f'{deal["DEAL_ID"].split("-")[-1]}')


def page_broker_note(deal: dict) -> list:
    broker = BROKERS[deal["BROKER_ID"]]
    cp = CPS[deal["COUNTERPARTY_ID"]]
    fx = deal["DEAL_TYPE"] in ("FX_SPOT", "FX_FORWARD")
    brokerage = float(deal["NOTIONAL_PKR"]) * 0.000015
    story = letterhead(broker, [
        f'{"Karachi" if deal["BROKER_ID"] != "BR-002" else "Lahore"}, Pakistan',
        "Member, Financial Markets Association of Pakistan"])
    story.append(Paragraph("BROKER'S CONTRACT NOTE", S_TITLE))
    story.append(kv([
        ("Note number", broker_note_ref(deal)),
        ("Date", long_date(deal["DEAL_DATE"])),
        ("To", f"{ENTITY}, Treasury"),
        ("We confirm having arranged the following transaction between you and",
         f'{cp["name"]}'),
        ("Product", PRODUCT[deal["DEAL_TYPE"]]),
        ("Principal", money(deal["NOTIONAL_AMOUNT"], deal["CURRENCY"])),
        ("Rate", f'{float(deal["DEAL_RATE"]):.4f}'
                 f'{"" if fx else " per cent per annum"}'),
        ("Value date", long_date(deal["VALUE_DATE"])),
        ("Maturity date", long_date(deal["MATURITY_DATE"])),
        ("Brokerage", money(round(brokerage, 2))),
    ]))
    story.append(Spacer(1, 14))
    story.append(signatures([("For and on behalf of the broker",
                              "Authorised signatory",
                              f'Dated {long_date(deal["DEAL_DATE"])}')]))
    return story


def page_confirmation(deal: dict, conf: dict, flags: set[str]) -> list:
    cp_id = deal["COUNTERPARTY_ID"]
    if "confirm_counterparty" in flags:
        cp_id = "CP-009"
    cp = CPS[cp_id]
    fx = deal["DEAL_TYPE"] in ("FX_SPOT", "FX_FORWARD")

    rate = float(deal["DEAL_RATE"])
    if "confirm_rate" in flags:
        rate = float(EXTRA["C2_confirmed_rate"])
    principal = float(deal["NOTIONAL_AMOUNT"])
    if "confirm_notional" in flags:
        principal = float(EXTRA["C3_confirmed_notional"])

    if "internal_confirmation" in flags:
        story = letterhead(ENTITY, ["Treasury operations, head office, Karachi",
                                    "Internal deal confirmation print"])
        story.append(Paragraph("DEAL CONFIRMATION", S_TITLE))
        story.append(kv([
            ("Printed by", person(conf["CONFIRMED_BY_ID"])),
            ("Printed on", long_date(conf["SENT_DATE"])),
            ("Deal reference", deal["DEAL_ID"]),
            ("Counterparty", f'{cp["name"]} ({cp_id})'),
            ("Product", PRODUCT[deal["DEAL_TYPE"]]),
            ("Principal", money(principal, deal["CURRENCY"])),
            ("Rate", f"{rate:.4f}{'' if fx else ' per cent per annum'}"),
            ("Value date", long_date(deal["VALUE_DATE"])),
            ("Maturity date", long_date(deal["MATURITY_DATE"])),
            ("Status", conf["MATCH_STATUS"]),
        ]))
        story.append(Spacer(1, 12))
        story.append(Paragraph(
            "Generated from the treasury system deal record. This print is for "
            "file purposes.", S_SMALL))
        story.append(Spacer(1, 14))
        story.append(signatures([("Prepared by", person(conf["CONFIRMED_BY_ID"]),
                                  f'Dated {long_date(conf["SENT_DATE"])}')]))
        return story

    story = letterhead(cp["name"], [
        f"Treasury operations, {CP_CITY[cp_id]}, Pakistan",
        f"SWIFT {CP_SWIFT[cp_id]}"])
    story.append(Paragraph(CONF_TITLE[conf["CONFIRMATION_MODE"]], S_TITLE))
    our_ref = (f'{cp["name"].split()[0][:3].upper()}/TR/CFM/'
               f'{deal["DEAL_DATE"][:4]}/{deal["DEAL_ID"].split("-")[-1]}')
    swift = [
        f':20:  {our_ref}',
        f':21:  {deal["DEAL_ID"]}',
        f':22A: NEWT',
        f':82A: {ENTITY_SWIFT}  {ENTITY}',
        f':87A: {CP_SWIFT[cp_id]}  {cp["name"]}',
        f':30T: {deal["DEAL_DATE"]}   trade date',
        f':30V: {deal["VALUE_DATE"]}   value date',
        f':30P: {deal["MATURITY_DATE"]}   maturity date',
        f':32B: {deal["CURRENCY"]} {principal:,.2f}',
        f':37G: {rate:.4f}{"" if fx else " per cent per annum"}',
        f':34E: {deal["CURRENCY"]}  settlement per standing instructions',
    ]
    story.append(kv([
        ("Our reference", our_ref),
        ("Your reference", deal["DEAL_ID"]),
        ("Date of issue", long_date(conf["SENT_DATE"])),
        ("Received by counterparty", long_date(conf["COUNTERPARTY_RECEIVED_DATE"])),
        ("Message type", conf["CONFIRMATION_MODE"]),
    ]))
    story.append(Spacer(1, 10))
    story.append(Paragraph(
        f"We confirm the following transaction concluded with your dealing "
        f"room on {long_date(deal['DEAL_DATE'])}.", S_BODY))
    story.append(Spacer(1, 8))
    for line in swift:
        story.append(Paragraph(line.replace(" ", "&nbsp;"), S_MONO))
    story.append(Spacer(1, 12))
    if conf["MATCH_STATUS"] == "Discrepancy" and conf["DISCREPANCY_NOTE"]:
        story.append(Paragraph(
            f'<b>Note added by {ENTITY} treasury operations on receipt:</b> '
            f'{conf["DISCREPANCY_NOTE"]}', S_BODY))
        story.append(Spacer(1, 10))
    story.append(Paragraph(
        "Please advise immediately if any of the above is at variance with "
        "your records.", S_BODY))
    story.append(Spacer(1, 14))
    story.append(signatures([
        ("For and on behalf of the counterparty", "Authorised signatory",
         f'Dated {long_date(conf["SENT_DATE"])}'),
        ("Matched by (Meridian treasury operations)",
         person(conf["CONFIRMED_BY_ID"]),
         f'Dated {long_date(conf["COUNTERPARTY_RECEIVED_DATE"])}')]))
    return story


def page_payment_instruction(deal: dict, stl: dict, flags: set[str]) -> list:
    story = letterhead(ENTITY, ["Treasury operations, head office, Karachi",
                                "Settlement section"])
    story.append(Paragraph("PAYMENT INSTRUCTION", S_TITLE))
    amount_pkr = float(stl["SETTLEMENT_AMOUNT_PKR"])
    story.append(kv([
        ("Instruction reference", stl["PAYMENT_REFERENCE"]),
        ("Settlement reference", stl["SETTLEMENT_ID"]),
        ("Deal reference", deal["DEAL_ID"]),
        ("Value date", long_date(stl["SETTLEMENT_DATE"])),
        ("Pay", money(stl["SETTLEMENT_AMOUNT"], stl["SETTLEMENT_CURRENCY"])),
        ("PKR equivalent", money(amount_pkr)),
        ("Beneficiary", CPS[deal["COUNTERPARTY_ID"]]["name"]),
        ("Beneficiary bank", stl["BENEFICIARY_BANK"]),
        ("Beneficiary account", stl["BENEFICIARY_ACCOUNT"]),
        ("Standing instruction", stl["SSI_ID"] or "None quoted"),
        ("Debit account", stl["NOSTRO_ACCOUNT"]),
    ]))
    story.append(Spacer(1, 14))
    blocks = [("Released by", person(stl["RELEASED_BY_ID"]),
               f'Dated {long_date(stl["SETTLEMENT_DATE"])}'),
              ("Approved by", person(stl["APPROVED_BY_ID"]),
               f'Dated {long_date(stl["SETTLEMENT_DATE"])}')]
    if amount_pkr >= 250_000_000 and "single_signature" not in flags:
        blocks.append(("Second approver (clause 8.5)",
                       person(stl["SECOND_APPROVER_ID"] or "TS-019"),
                       f'Dated {long_date(stl["SETTLEMENT_DATE"])}'))
    story.append(signatures(blocks))
    story.append(Spacer(1, 10))
    if amount_pkr >= 250_000_000:
        story.append(Paragraph(
            "Instructions for PKR 250,000,000 or more require two authorised "
            "signatures on this instruction in addition to system release and "
            "approval (Treasury and Investment Policy, clause 8.5).", S_SMALL))
    return story


def page_nostro(deal: dict, stl: dict) -> list:
    account = stl["NOSTRO_ACCOUNT"]
    currency = stl["SETTLEMENT_CURRENCY"]
    amount = float(stl["SETTLEMENT_AMOUNT"])
    opening = round(amount * 4.7 + 91_450_000, 2)
    story = letterhead(ENTITY, [f"Account statement extract - {account}",
                                f"Currency {currency}"])
    story.append(Paragraph("STATEMENT EXTRACT", S_TITLE))
    settle = date.fromisoformat(stl["SETTLEMENT_DATE"])
    rows = [["Date", "Narrative", "Debit", "Credit", "Balance"]]
    balance = opening
    filler = [
        (add_bd(settle, -0) - timedelta(days=1),
         "Interbank clearing - net", 0.0, round(amount * 0.31, 2)),
    ]
    for when, narrative, debit, credit in filler:
        balance = balance - debit + credit
        rows.append([when.isoformat(), narrative,
                     f"{debit:,.2f}" if debit else "",
                     f"{credit:,.2f}" if credit else "", f"{balance:,.2f}"])
    balance -= amount
    rows.append([stl["SETTLEMENT_DATE"],
                 f'{stl["PAYMENT_REFERENCE"]} {CPS[deal["COUNTERPARTY_ID"]]["name"]}'
                 f' - {deal["DEAL_ID"]}',
                 f"{amount:,.2f}", "", f"{balance:,.2f}"])
    credit_back = round(amount * 0.12, 2)
    balance += credit_back
    rows.append([add_bd(settle, 1).isoformat(), "Interbank clearing - net", "",
                 f"{credit_back:,.2f}", f"{balance:,.2f}"])

    table = Table(rows, colWidths=(22 * mm, 76 * mm, 24 * mm, 24 * mm, 26 * mm))
    table.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", 8),
        ("FONT", (0, 1), (-1, -1), "Helvetica", 8),
        ("ALIGN", (2, 0), (-1, -1), "RIGHT"),
        ("LINEBELOW", (0, 0), (-1, 0), 0.6, colors.HexColor("#222222")),
        ("LINEBELOW", (0, 1), (-1, -2), 0.25, colors.HexColor("#DDDDDD")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(table)
    story.append(Spacer(1, 10))
    story.append(Paragraph(
        f"Extract produced by the nostro reconciliation officer for the "
        f"internal audit sample. Opening balance {money(opening, currency)} "
        f"brought forward.", S_SMALL))
    return story


# --------------------------------------------------------------------------
# pack assembly
# --------------------------------------------------------------------------

FLAGS = {
    "D1": {"no_authorisation"},
    "D2": {"altered_rate"},
    "D3": {"internal_confirmation"},
    "D4": {"no_broker_note"},
    "D5": {"single_signature"},
    "D6": {"late_authorisation"},
    "C1": {"confirm_counterparty"},
    "C2": {"confirm_rate"},
    "C3": {"confirm_notional"},
    "C4": {"synthetic"},
}


def footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 7)
    canvas.setFillColor(colors.HexColor("#888888"))
    canvas.drawString(20 * mm, 12 * mm,
                      f"{ENTITY} - internal audit sample. Fictional document.")
    canvas.drawRightString(A4[0] - 20 * mm, 12 * mm, f"Page {doc.page}")
    canvas.restoreState()


def emit(path: Path, story: list, title: str) -> None:
    doc = SimpleDocTemplate(
        str(path), pagesize=A4,
        leftMargin=21 * mm, rightMargin=21 * mm,
        topMargin=18 * mm, bottomMargin=20 * mm,
        title=title,
        author=f"{ENTITY} (fictional)",
        subject="Internal audit sample - treasury dealing")
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def build_pack(ordinal: int, code: str, deal_id: str) -> list[Path]:
    """One folder per sampled deal, one PDF per document inside it.

    Intake treats a file as a document, so the ticket, the broker note, the
    confirmation, the instruction and the statement extract cannot share a PDF.
    Filenames carry the document's own reference and its type, because intake
    classification reads the filename and not the contents.
    """
    flags = FLAGS.get(code, set())
    if "synthetic" in flags:
        deal, conf, stl = SYNTHETIC, SYNTHETIC_CONF, SYNTHETIC_STL
    else:
        deal = DEALS[deal_id]
        conf = CONFS.get(deal_id)
        stl = SETTS.get(deal_id)
    if conf is None or stl is None:
        raise SystemExit(f"{code} {deal_id}: pack needs both a confirmation "
                         f"and a settlement")

    folder = PACKS_DIR / f"{ordinal:02d}_{deal['DEAL_ID']}"
    folder.mkdir(parents=True, exist_ok=True)

    documents = [
        (f"{deal['DEAL_ID']}_Dealing_Ticket.pdf",
         page_deal_ticket(deal, flags),
         f"Dealing ticket {deal['DEAL_ID']}"),
    ]
    if deal["BROKER_ID"] and "no_broker_note" not in flags:
        documents.append(
            (f"{broker_note_ref(deal)}_Broker_Note.pdf",
             page_broker_note(deal),
             f"Broker note {broker_note_ref(deal)}"))
    documents += [
        (f"{conf['CONFIRMATION_ID']}_Counterparty_Confirmation.pdf",
         page_confirmation(deal, conf, flags),
         f"Counterparty confirmation {conf['CONFIRMATION_ID']}"),
        (f"{stl['PAYMENT_REFERENCE']}_Payment_Instruction.pdf",
         page_payment_instruction(deal, stl, flags),
         f"Payment instruction {stl['PAYMENT_REFERENCE']}"),
        (f"{stl['SETTLEMENT_ID']}_Nostro_Statement_Extract.pdf",
         page_nostro(deal, stl),
         f"Nostro statement extract {stl['SETTLEMENT_ID']}"),
    ]

    written = []
    for name, story, title in documents:
        path = folder / name
        emit(path, story, title)
        written.append(path)
    return written


if __name__ == "__main__":
    ROOT.mkdir(parents=True, exist_ok=True)
    PACKS_DIR.mkdir(parents=True, exist_ok=True)
    build_policy()
    build_limits()
    build_minutes()
    print("criteria documents written")
    total = 0
    for ordinal, (code, deal_id) in enumerate(
            sorted(PACKS.items(), key=lambda kv: kv[1]), 1):
        written = build_pack(ordinal, code, deal_id)
        total += len(written)
        print(f"  {code:5} {written[0].parent.name}: "
              f"{len(written)} documents")
    print(f"{len(PACKS)} deal folders, {total} documents written")
