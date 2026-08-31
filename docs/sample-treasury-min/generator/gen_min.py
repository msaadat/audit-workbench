"""Derive the minimised treasury pack from docs/sample-treasury.

The minimised pack is a workflow-mechanism fixture, not an audit-quality
fixture: one policy document, two tables and one deal pack, sized so a full
intake -> RCM -> analysis -> vouch cycle runs in a fraction of the time the
full pack takes. Every row and every PDF is copied unaltered from the full
pack, so the two agree wherever they overlap.

Run from anywhere:  python docs/sample-treasury-min/generator/gen_min.py
"""

from __future__ import annotations

import csv
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

MIN_ROOT = Path(__file__).resolve().parent.parent
FULL_ROOT = MIN_ROOT.parent / "sample-treasury"

# The deal window. Contiguous, so no spurious reference-sequence gaps.
FIRST_DEAL = "TD-2025-0160"
LAST_DEAL = "TD-2025-0200"

# The one sampled deal, and the folder it is staged under.
PACK_DEAL = "TD-2025-0166"
PACK_SRC = FULL_ROOT / "documents" / "deal-packs" / "08_TD-2025-0166"
PACK_DST = MIN_ROOT / "documents" / "deal-packs" / "01_TD-2025-0166"

ENTITY = "Meridian Bank Limited"

# --------------------------------------------------------------------------
# edits to the copied rows
#
# The window is copied verbatim except for the changes below. Their purpose is
# to leave exactly one seeded exception per control, and to remove incidental
# breaks that would be raised as findings without being seeded as any.
# --------------------------------------------------------------------------

# NOTIONAL_PKR is a half-rupee short of NOTIONAL_AMOUNT x DEAL_RATE in the full
# pack. Harmless there; here it is the only recomputation break in 41 rows, so
# it would read as the seeded exception it is not.
DEAL_EDITS = {
    "TD-2025-0196": {"NOTIONAL_PKR": "51095082.50"},
}

# The full pack seeds the same segregation failure twice (TD-2025-0167 and
# TD-2025-0168). One is enough to prove the control fires, so TD-2025-0168 is
# reassigned to a back-office confirmer.
CONF_EDITS = {
    "CNF-2025-0168": {"CONFIRMED_BY_ID": "TS-016"},
}

# The full pack leaves three settled deals unconfirmed (TD-2025-0180, 0181 and
# 0183). TD-2025-0180 stays unconfirmed; the other two get the confirmation
# they should have had.
CONF_INSERTS = [
    {"CONFIRMATION_ID": "CNF-2025-0181", "DEAL_ID": "TD-2025-0181",
     "CONFIRMATION_MODE": "SWIFT MT300", "SENT_DATE": "2025-02-06",
     "COUNTERPARTY_RECEIVED_DATE": "2025-02-07", "CONFIRMED_BY_ID": "TS-020",
     "MATCH_STATUS": "Matched", "DISCREPANCY_NOTE": ""},
    {"CONFIRMATION_ID": "CNF-2025-0183", "DEAL_ID": "TD-2025-0183",
     "CONFIRMATION_MODE": "SWIFT MT300", "SENT_DATE": "2025-02-07",
     "COUNTERPARTY_RECEIVED_DATE": "2025-02-07", "CONFIRMED_BY_ID": "TS-016",
     "MATCH_STATUS": "Matched", "DISCREPANCY_NOTE": ""},
]


def subset_csv(src: Path, dst: Path, key: str, id_col: str,
               edits: dict | None = None,
               inserts: list[dict] | None = None) -> int:
    with src.open(newline="", encoding="utf-8") as fh:
        rows = [r for r in csv.DictReader(fh) if FIRST_DEAL <= r[key] <= LAST_DEAL]
        fh.seek(0)
        header = next(csv.reader(fh))
    for row in rows:
        row.update(edits.get(row[id_col], {}) if edits else {})
    rows.extend(inserts or [])
    rows.sort(key=lambda r: r[id_col])
    dst.parent.mkdir(parents=True, exist_ok=True)
    with dst.open("w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=header)
        writer.writeheader()
        writer.writerows(rows)
    return len(rows)


def build_policy() -> None:
    """The policy extract, cut to the clauses the two tables can be tested against.

    Sections 6 and 8, and the limit clauses of 4 and 5, are dropped: there is no
    market rate file, no settlement file and no limit file in this pack, so
    carrying those clauses would only author controls that cannot be tested.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    title = doc.add_paragraph()
    run = title.add_run(f"{ENTITY.upper()}\nTREASURY AND INVESTMENT POLICY")
    run.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.add_run("Extract for internal audit - dealing capture and confirmation\n").italic = True
    sub.add_run("Document reference TP/2023/04  |  Version 4.1  |  Approved by "
                "the Board Risk Committee on 12 October 2023  |  Effective 1 "
                "November 2023  |  Next review due 1 November 2024").italic = True
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    sections = [
        ("4. Organisation and segregation of duties", [
            ("4.1", "Dealing, confirmation and settlement are performed by "
             "separate functions. A member of the front office shall not "
             "confirm, verify or release settlement of any deal, whether or "
             "not that deal is one they executed."),
        ]),
        ("5. Dealing", [
            ("5.2", "Dealing is permitted between 09:00 and 17:00 on business "
             "days only. Any deal struck outside those hours, or on a "
             "non-business day, requires the prior written approval of the "
             "Head of Treasury and must be reported to ALCO at its next "
             "meeting."),
            ("5.4", "Every deal is captured in the treasury system within 15 "
             "minutes of execution. Capture records the dealer's identity and "
             "the time of execution as distinct from the time of capture."),
            ("5.6", "A deal may not be amended or cancelled after capture "
             "without the approval of the Head of Treasury Operations. The "
             "approver's identity and the date of approval are recorded "
             "against the deal."),
        ]),
        ("7. Confirmation", [
            ("7.2", "Confirmations are despatched no later than one business "
             "day after the deal date."),
            ("7.3", "The counterparty's confirmation is matched to the deal "
             "record on counterparty, principal, rate, value date and maturity "
             "date. A discrepancy is referred to the Head of Treasury "
             "Operations and is resolved before settlement."),
            ("7.4", "No deal is settled before the counterparty's confirmation "
             "has been received and matched. A confirmation produced "
             "internally is not evidence of the counterparty's agreement."),
        ]),
    ]
    for name, clauses in sections:
        doc.add_heading(name, level=1)
        for number, text in clauses:
            para = doc.add_paragraph()
            para.add_run(f"{number}  ").bold = True
            para.add_run(text)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("This extract is issued for the purposes of the internal audit "
                 "of treasury dealing capture and confirmation for the half "
                 "year ended 30 June 2025. It reproduces the clauses of "
                 "sections 4, 5 and 7 that bear on deal capture and "
                 "counterparty confirmation. The counterparty and dealer limit "
                 "clauses, section 6 (rates and valuation) and section 8 "
                 "(settlement) are not reproduced and are outside this scope."
                 ).italic = True

    out = MIN_ROOT / "documents" / "01_Treasury_and_Investment_Policy.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)


def copy_pack() -> int:
    if PACK_DST.exists():
        shutil.rmtree(PACK_DST)
    shutil.copytree(PACK_SRC, PACK_DST)
    return len(list(PACK_DST.glob("*.pdf")))


def main() -> None:
    deals = subset_csv(FULL_ROOT / "data" / "04_deals.csv",
                       MIN_ROOT / "data" / "04_deals.csv",
                       "DEAL_ID", "DEAL_ID", edits=DEAL_EDITS)
    confs = subset_csv(FULL_ROOT / "data" / "05_confirmations.csv",
                       MIN_ROOT / "data" / "05_confirmations.csv",
                       "DEAL_ID", "CONFIRMATION_ID",
                       edits=CONF_EDITS, inserts=CONF_INSERTS)
    build_policy()
    pdfs = copy_pack()
    print(f"deals {deals} rows, confirmations {confs} rows, "
          f"1 policy docx, 1 deal pack of {pdfs} PDFs")


if __name__ == "__main__":
    main()
